"""Generate PDF, TXT, and DOCX fixtures used by feedback-ingestion tests.

Re-run whenever fixtures need to be rebuilt:
    python backend/scripts/build_pdf_test_fixtures.py

Outputs into ``backend/tests/fixtures/``:

* ``mock_feedback.pdf`` — multi-page PDF with three paragraphs separated by blank lines.
* ``mock_feedback_scanned.pdf`` — single page containing only an image (no
  extractable text), used to assert we surface a clear error.
* ``mock_feedback_encrypted.pdf`` — password-protected PDF.
* ``mock_feedback.txt`` — text file with mixed line endings, a UTF-8 BOM, and
  a non-ASCII (Tamil) line.
* ``mock_feedback.docx`` — Word document with three non-empty paragraphs
  interspersed with blank paragraphs (which must be filtered out).
"""

from __future__ import annotations

import io
from pathlib import Path

from PIL import Image
from docx import Document
from pypdf import PdfReader, PdfWriter
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas

FIXTURES_DIR = Path(__file__).resolve().parents[1] / "tests" / "fixtures"
PARAGRAPHS = [
    "The new dashboard layout is wonderful and feels much faster on my laptop.",
    "However the export button keeps failing on Safari with a generic error message.",
    "I would love to see better keyboard shortcuts for the comment review queue.",
]


def build_text_pdf(target: Path) -> None:
    """Write a 2-page PDF that contains PARAGRAPHS separated by blank lines."""
    c = canvas.Canvas(str(target), pagesize=LETTER)
    width, height = LETTER

    # Page 1 — first two paragraphs separated by a blank line.
    text_object = c.beginText(1 * inch, height - 1 * inch)
    text_object.setFont("Helvetica", 12)
    text_object.textLine(PARAGRAPHS[0])
    text_object.textLine("")
    text_object.textLine(PARAGRAPHS[1])
    c.drawText(text_object)
    c.showPage()

    # Page 2 — third paragraph.
    text_object = c.beginText(1 * inch, height - 1 * inch)
    text_object.setFont("Helvetica", 12)
    text_object.textLine(PARAGRAPHS[2])
    c.drawText(text_object)
    c.showPage()
    c.save()


def build_scanned_pdf(target: Path) -> None:
    """Write a single-page PDF whose only content is a JPEG image (no text)."""
    image = Image.new("RGB", (400, 200), color=(220, 220, 220))
    image_buf = io.BytesIO()
    image.save(image_buf, format="JPEG")
    image_buf.seek(0)

    c = canvas.Canvas(str(target), pagesize=LETTER)
    width, height = LETTER
    c.drawImage(
        image=__import__("reportlab.lib.utils", fromlist=["ImageReader"]).ImageReader(image_buf),
        x=1 * inch,
        y=height - 4 * inch,
        width=4 * inch,
        height=2 * inch,
    )
    c.showPage()
    c.save()


def build_encrypted_pdf(source: Path, target: Path, password: str = "saramsa") -> None:
    """Encrypt an existing PDF in place into ``target``."""
    reader = PdfReader(str(source))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt(user_password=password, owner_password=password)
    with target.open("wb") as fh:
        writer.write(fh)


def build_docx_fixture(target: Path) -> None:
    """Write a Word document with three non-empty paragraphs and a few blanks."""
    doc = Document()
    doc.add_paragraph(PARAGRAPHS[0])
    doc.add_paragraph("")  # blank paragraph — must be filtered
    doc.add_paragraph(PARAGRAPHS[1])
    doc.add_paragraph("   ")  # whitespace-only — must be filtered
    doc.add_paragraph(PARAGRAPHS[2])
    doc.save(str(target))


def build_text_fixture(target: Path) -> None:
    """Write a text fixture with mixed line endings, a BOM, and Tamil text."""
    body = (
        "﻿"  # BOM
        "First feedback comment\r\n"
        "\r\n"  # blank line — should be filtered
        "Second feedback comment\r"
        "மூன்றாவது கருத்து\n"  # "Third comment" in Tamil
    )
    target.write_bytes(body.encode("utf-8"))


def main() -> None:
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)

    text_pdf = FIXTURES_DIR / "mock_feedback.pdf"
    scanned_pdf = FIXTURES_DIR / "mock_feedback_scanned.pdf"
    encrypted_pdf = FIXTURES_DIR / "mock_feedback_encrypted.pdf"
    text_file = FIXTURES_DIR / "mock_feedback.txt"
    docx_file = FIXTURES_DIR / "mock_feedback.docx"

    build_text_pdf(text_pdf)
    build_scanned_pdf(scanned_pdf)
    build_encrypted_pdf(text_pdf, encrypted_pdf)
    build_text_fixture(text_file)
    build_docx_fixture(docx_file)

    print(f"Wrote fixtures to {FIXTURES_DIR}")


if __name__ == "__main__":
    main()
