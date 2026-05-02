"""Convert uploaded PDF, plain-text, and Word (.docx) files into a list of
feedback comments.

Each extractor returns ``list[str]`` and raises ``ValueError`` with a
user-friendly message when the file cannot be turned into comments. The
caller (``FeedbackFileIngestView``) translates that into an HTTP 400.

These are pure functions and intentionally have no Django dependency so they
can be unit-tested without bootstrapping the wider feedback-analysis stack.
"""

from __future__ import annotations

import logging
import zipfile
from typing import IO, List

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from pypdf import PdfReader
from pypdf.errors import FileNotDecryptedError, PdfReadError

logger = logging.getLogger(__name__)

# Reject DOCX whose inner ZIP would decompress to more than this many bytes.
# Defends against zip-bomb / decompression-DoS attacks where a small upload
# explodes into gigabytes of XML once python-docx + lxml parse it.
_DOCX_MAX_UNCOMPRESSED_BYTES = 50 * 1024 * 1024  # 50 MB
_DOCX_STREAM_CHUNK_BYTES = 64 * 1024


def _enforce_docx_size_cap(file_obj: IO[bytes]) -> None:
    """Stream-decompress every zip member; raise if the cumulative
    decompressed bytes exceed the cap. Leaves ``file_obj`` rewound for
    the caller (python-docx) to re-parse."""
    try:
        with zipfile.ZipFile(file_obj) as zf:
            total = 0
            for zi in zf.infolist():
                with zf.open(zi) as member:
                    while True:
                        chunk = member.read(_DOCX_STREAM_CHUNK_BYTES)
                        if not chunk:
                            break
                        total += len(chunk)
                        if total > _DOCX_MAX_UNCOMPRESSED_BYTES:
                            raise ValueError(
                                "DOCX is too large to process. "
                                "Please split it into smaller files."
                            )
    except zipfile.BadZipFile as exc:
        logger.warning("DOCX is not a valid zip archive: %s", exc)
        raise ValueError(
            "DOCX could not be read. The file may be corrupted or password-protected."
        ) from exc
    file_obj.seek(0)

_BOM = "﻿"


def _split_lines_to_comments(text: str) -> List[str]:
    """Normalize line endings, strip whitespace, and drop empty entries."""
    if not text:
        return []
    if text.startswith(_BOM):
        text = text[len(_BOM):]
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    return [line.strip() for line in normalized.split("\n") if line.strip()]


def extract_comments_from_text(file_obj: IO[bytes]) -> List[str]:
    """Read a UTF-8 text upload and return one comment per non-empty line."""
    raw = file_obj.read()
    decoded = raw.decode("utf-8", errors="replace")
    comments = _split_lines_to_comments(decoded)
    if not comments:
        raise ValueError("Text file is empty.")
    return comments


def extract_comments_from_pdf(file_obj: IO[bytes]) -> List[str]:
    """Read a PDF upload and return one comment per non-empty extracted line."""
    try:
        reader = PdfReader(file_obj)
    except PdfReadError as exc:
        raise ValueError(f"PDF could not be read: {exc}") from exc

    if reader.is_encrypted:
        raise ValueError(
            "PDF is encrypted. Please remove the password before uploading."
        )

    comments: List[str] = []
    try:
        for page in reader.pages:
            page_text = page.extract_text() or ""
            comments.extend(_split_lines_to_comments(page_text))
    except FileNotDecryptedError as exc:
        raise ValueError(
            "PDF is encrypted. Please remove the password before uploading."
        ) from exc
    except PdfReadError as exc:
        raise ValueError(f"PDF could not be read: {exc}") from exc

    if not comments:
        raise ValueError(
            "PDF contains no extractable text. "
            "Encrypted or image-only (scanned) PDFs are not supported."
        )
    return comments


def _append_lines(comments: List[str], text: str) -> None:
    """Append non-empty lines from ``text`` (soft line breaks render as
    ``\\n`` inside a paragraph). Delegates to ``_split_lines_to_comments``
    so DOCX ends up with the same BOM-strip / CR-normalize / trim
    semantics as TXT and PDF — the format-parity claim in the docs
    relies on it."""
    comments.extend(_split_lines_to_comments(text))


def _walk_tables(tables, comments: List[str]) -> None:
    """Recursively collect text from a list of python-docx tables, including
    arbitrarily nested sub-tables."""
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _append_lines(comments, paragraph.text)
                # getattr fallback for python-docx versions that pre-date
                # _Cell.tables (we pin >=1.2.0 but stay defensive).
                nested = getattr(cell, "tables", None) or []
                if nested:
                    _walk_tables(nested, comments)


def extract_comments_from_docx(file_obj: IO[bytes]) -> List[str]:
    """Read a Word (.docx) upload and return one comment per non-empty line.

    Each paragraph is treated as a comment boundary, and soft line breaks
    (Shift+Enter / ``<w:br/>``) inside a paragraph are split into separate
    comments — matching the per-line semantics of PDF/TXT. Text inside
    tables (including nested tables) is collected too, since users
    sometimes paste feedback as a single-column table.

    ``file_obj`` must be a seekable byte stream; the size-cap pass rewinds
    it before handing to python-docx. ``request.FILES`` uploads from
    Django's MemoryFileUploadHandler / TemporaryFileUploadHandler are
    seekable, as is ``io.BytesIO``.
    """
    # Stream-decompress every member with a hard byte ceiling. Cannot trust
    # ZipInfo.file_size — a zip-bomb crafts the central directory to claim
    # a tiny size while the actual stream expands to gigabytes. The only
    # safe defense is to bound the bytes we actually decompress.
    _enforce_docx_size_cap(file_obj)

    try:
        doc = Document(file_obj)
    except (PackageNotFoundError, zipfile.BadZipFile) as exc:
        logger.warning("DOCX failed to open via python-docx: %s", exc)
        raise ValueError(
            "DOCX could not be read. The file may be corrupted or password-protected."
        ) from exc

    comments: List[str] = []
    for paragraph in doc.paragraphs:
        _append_lines(comments, paragraph.text)
    _walk_tables(doc.tables, comments)

    if not comments:
        raise ValueError(
            "DOCX contains no extractable text. The document appears to be empty."
        )
    return comments
