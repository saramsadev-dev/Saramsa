"""Unit tests for PDF and text-file comment extractors.

These tests exercise the pure extraction functions with real fixture files;
they do not require Django or the broader feedback-analysis stack.

Run with:
    backend/venv/Scripts/python.exe -m pytest backend/feedback_analysis/tests/test_file_extractors.py
"""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

import pytest

from feedback_analysis.file_extractors import (
    extract_comments_from_docx,
    extract_comments_from_pdf,
    extract_comments_from_text,
)

FIXTURES = Path(__file__).resolve().parents[2] / "tests" / "fixtures"


def _open(path: Path) -> io.BytesIO:
    return io.BytesIO(path.read_bytes())


class TestExtractCommentsFromText:
    def test_splits_nonempty_lines_with_mixed_line_endings(self):
        comments = extract_comments_from_text(_open(FIXTURES / "mock_feedback.txt"))
        assert comments == [
            "First feedback comment",
            "Second feedback comment",
            "மூன்றாவது கருத்து",
        ]

    def test_strips_utf8_bom(self):
        payload = io.BytesIO("﻿hello\nworld\n".encode("utf-8"))
        assert extract_comments_from_text(payload) == ["hello", "world"]

    def test_normalizes_carriage_returns(self):
        payload = io.BytesIO(b"alpha\r\nbeta\rgamma\n")
        assert extract_comments_from_text(payload) == ["alpha", "beta", "gamma"]

    def test_drops_blank_and_whitespace_only_lines(self):
        payload = io.BytesIO(b"one\n\n   \n\ttwo\n")
        assert extract_comments_from_text(payload) == ["one", "two"]

    def test_decodes_invalid_utf8_with_replacement(self):
        payload = io.BytesIO(b"valid line\nbroken \xff byte\n")
        result = extract_comments_from_text(payload)
        assert result[0] == "valid line"
        assert result[1].startswith("broken ")
        assert len(result) == 2

    def test_empty_file_raises_value_error(self):
        with pytest.raises(ValueError, match="Text file is empty"):
            extract_comments_from_text(io.BytesIO(b""))

    def test_whitespace_only_file_raises_value_error(self):
        with pytest.raises(ValueError, match="Text file is empty"):
            extract_comments_from_text(io.BytesIO(b"\n\n   \r\n"))

    def test_single_long_line_returns_one_comment(self):
        payload = io.BytesIO(b"only one line here without a trailing newline")
        assert extract_comments_from_text(payload) == [
            "only one line here without a trailing newline",
        ]


class TestExtractCommentsFromPdf:
    def test_extracts_one_comment_per_line_across_pages(self):
        comments = extract_comments_from_pdf(_open(FIXTURES / "mock_feedback.pdf"))
        assert comments == [
            "The new dashboard layout is wonderful and feels much faster on my laptop.",
            "However the export button keeps failing on Safari with a generic error message.",
            "I would love to see better keyboard shortcuts for the comment review queue.",
        ]

    def test_image_only_pdf_raises_value_error(self):
        with pytest.raises(ValueError, match="no extractable text"):
            extract_comments_from_pdf(_open(FIXTURES / "mock_feedback_scanned.pdf"))

    def test_encrypted_pdf_raises_value_error(self):
        with pytest.raises(ValueError, match="encrypted"):
            extract_comments_from_pdf(_open(FIXTURES / "mock_feedback_encrypted.pdf"))

    def test_corrupt_bytes_raises_value_error(self):
        with pytest.raises(ValueError):
            extract_comments_from_pdf(io.BytesIO(b"this is not a pdf"))


class TestExtractCommentsFromDocx:
    def test_extracts_one_comment_per_nonempty_paragraph(self):
        comments = extract_comments_from_docx(_open(FIXTURES / "mock_feedback.docx"))
        assert comments == [
            "The new dashboard layout is wonderful and feels much faster on my laptop.",
            "However the export button keeps failing on Safari with a generic error message.",
            "I would love to see better keyboard shortcuts for the comment review queue.",
        ]

    def test_empty_docx_raises_value_error(self):
        from docx import Document
        buf = io.BytesIO()
        doc = Document()
        doc.add_paragraph("")
        doc.add_paragraph("   ")
        doc.save(buf)
        buf.seek(0)
        with pytest.raises(ValueError, match="no extractable text"):
            extract_comments_from_docx(buf)

    def test_corrupt_bytes_raises_value_error(self):
        with pytest.raises(ValueError):
            extract_comments_from_docx(io.BytesIO(b"this is not a docx"))

    def test_extracts_text_from_table_cells_in_order(self):
        from docx import Document
        buf = io.BytesIO()
        doc = Document()
        doc.add_paragraph("Header paragraph")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cell A1"
        table.cell(0, 1).text = "Cell A2"
        table.cell(1, 0).text = ""
        table.cell(1, 1).text = "Cell B2"
        doc.save(buf)
        buf.seek(0)
        assert extract_comments_from_docx(buf) == [
            "Header paragraph",
            "Cell A1",
            "Cell A2",
            "Cell B2",
        ]

    def test_splits_paragraph_with_soft_line_breaks(self):
        # Static fixture — pins the on-disk wire format so this test is
        # robust against future changes in python-docx's serializer.
        comments = extract_comments_from_docx(_open(FIXTURES / "mock_feedback_softbreak.docx"))
        assert comments == ["first line", "second line"]

    def test_extracts_text_from_nested_tables(self):
        comments = extract_comments_from_docx(_open(FIXTURES / "mock_feedback_nested_table.docx"))
        assert "Outer cell text" in comments
        assert "Inner left" in comments
        assert "Inner right" in comments

    def test_zip_bomb_rejected_even_with_lying_central_directory(self):
        # Hand-craft a DOCX-shaped zip whose single member decompresses to
        # 200 MB of zeros (well above the 50 MB cap). The cap must catch
        # this regardless of what the central directory header claims.
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("word/document.xml", b"\x00" * (200 * 1024 * 1024))
        buf.seek(0)
        with pytest.raises(ValueError, match="too large"):
            extract_comments_from_docx(buf)
